"""
Procesador de documentos multimodal
TFM Vicente Caruncho - Sistemas Inteligentes
"""

import io
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime
from app.models.document import DocumentChunk, DocumentMetadata


# Imports para procesamiento de documentos
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from app.core.logger import get_logger
from app.core.config import get_rag_config


class DocumentProcessor:
    """Procesador de documentos multimodal"""
    
    def __init__(self):
        self.logger = get_logger("document_processor")
        
        # Configuración de chunking
        rag_config = get_rag_config()
        self.chunk_size = rag_config.chunk_size
        self.chunk_overlap = rag_config.chunk_overlap
        self.separator = getattr(rag_config, 'separator', '\n\n')
        
        # Mapeo de extensiones a métodos de procesamiento
        self.processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # Intentar con docx
            '.txt': self._process_text,
            '.md': self._process_text,
            '.rtf': self._process_text,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
        }
        
        self.logger.info(f"DocumentProcessor inicializado - Chunk size: {self.chunk_size}")
    
    def process_file(self, file_path: str) -> List[DocumentChunk]:
        """Procesar archivo y generar chunks"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.processors:
            raise ValueError(f"Extensión no soportada: {extension}")
        
        try:
            # Procesar según tipo de archivo
            text_content = self.processors[extension](file_path)
            
            if not text_content or not text_content.strip():
                self.logger.warning(f"No se extrajo contenido de: {path.name}")
                return []
            
            # Generar chunks
            chunks = self._create_chunks(text_content, file_path)
            
            self.logger.info(f"Procesado {path.name}: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error procesando {path.name}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> str:
        """Procesar archivo PDF"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 no está disponible")
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Página {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        self.logger.warning(f"Error extrayendo página {page_num + 1}: {e}")
                        continue
        
        except Exception as e:
            raise Exception(f"Error leyendo PDF: {e}")
        
        return '\n\n'.join(text_content)
    
    def _process_docx(self, file_path: str) -> str:
        """Procesar archivo DOCX/DOC"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está disponible")
        
        try:
            doc = DocxDocument(file_path)
            
            text_content = []
            
            # Extraer párrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
            
            # Extraer tablas
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_content.append('\n--- Tabla ---\n' + '\n'.join(table_text))
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Error leyendo DOCX: {e}")
    
    def _process_text(self, file_path: str) -> str:
        """Procesar archivo de texto plano"""
        try:
            # Intentar diferentes codificaciones
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("No se pudo decodificar el archivo con ninguna codificación")
            
        except Exception as e:
            raise Exception(f"Error leyendo archivo de texto: {e}")
    
    def _process_csv(self, file_path: str) -> str:
        """Procesar archivo CSV"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas no está disponible")
        
        try:
            # Leer CSV con diferentes configuraciones
            try:
                df = pd.read_csv(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, encoding='latin-1')
            
            # Convertir a texto estructurado
            text_content = [f"Archivo CSV: {Path(file_path).name}"]
            text_content.append(f"Columnas: {', '.join(df.columns.tolist())}")
            text_content.append(f"Filas: {len(df)}")
            text_content.append("\n--- Datos ---")
            
            # Incluir primeras filas como muestra
            sample_size = min(50, len(df))
            for idx, row in df.head(sample_size).iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                
                if row_text:
                    text_content.append(f"Fila {idx + 1}: {' | '.join(row_text)}")
            
            if len(df) > sample_size:
                text_content.append(f"\n... y {len(df) - sample_size} filas adicionales")
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Error leyendo CSV: {e}")
    
    def _process_excel(self, file_path: str) -> str:
        """Procesar archivo Excel"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas no está disponible")
        
        try:
            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            text_content = [f"Archivo Excel: {Path(file_path).name}"]
            text_content.append(f"Hojas: {', '.join(excel_file.sheet_names)}")
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    text_content.append(f"\n--- Hoja: {sheet_name} ---")
                    text_content.append(f"Columnas: {', '.join(df.columns.tolist())}")
                    text_content.append(f"Filas: {len(df)}")
                    
                    # Muestra de datos
                    sample_size = min(20, len(df))
                    for idx, row in df.head(sample_size).iterrows():
                        row_text = []
                        for col, value in row.items():
                            if pd.notna(value):
                                row_text.append(f"{col}: {value}")
                        
                        if row_text:
                            text_content.append(f"  {' | '.join(row_text)}")
                    
                    if len(df) > sample_size:
                        text_content.append(f"  ... y {len(df) - sample_size} filas adicionales")
                        
                except Exception as e:
                    text_content.append(f"Error procesando hoja {sheet_name}: {e}")
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Error leyendo Excel: {e}")
    
    def _create_chunks(self, text: str, source_file: str) -> List[DocumentChunk]:
        """Crear chunks de texto usando estrategia configurable"""
        
        if not text.strip():
            return []
        
        chunks = []
        
        # Dividir por separador principal
        sections = text.split(self.separator)
        
        current_chunk = ""
        chunk_index = 0
        start_char = 0
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            # Si la sección cabe en el chunk actual
            test_chunk = f"{current_chunk}{self.separator}{section}" if current_chunk else section
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                # Guardar chunk actual si no está vacío
                if current_chunk:
                    chunk = self._create_chunk(
                        content=current_chunk,
                        source_file=source_file,
                        chunk_index=chunk_index,
                        start_char=start_char,
                        end_char=start_char + len(current_chunk)
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Manejar secciones muy largas
                if len(section) > self.chunk_size:
                    # Dividir sección larga en chunks más pequeños
                    sub_chunks = self._split_long_section(
                        section, source_file, chunk_index, start_char + len(current_chunk)
                    )
                    chunks.extend(sub_chunks)
                    chunk_index += len(sub_chunks)
                    
                    # Reiniciar con overlap si hay chunks anteriores
                    if chunks:
                        overlap_text = section[-self.chunk_overlap:] if len(section) > self.chunk_overlap else section
                        current_chunk = overlap_text
                    else:
                        current_chunk = ""
                else:
                    # Aplicar overlap del chunk anterior
                    if chunks and self.chunk_overlap > 0:
                        previous_chunk = chunks[-1].content
                        overlap_text = previous_chunk[-self.chunk_overlap:] if len(previous_chunk) > self.chunk_overlap else previous_chunk
                        current_chunk = f"{overlap_text}{self.separator}{section}"
                    else:
                        current_chunk = section
                
                start_char = start_char + len(current_chunk) - (self.chunk_overlap if chunks else 0)
        
        # Agregar último chunk si no está vacío
        if current_chunk.strip():
            chunk = self._create_chunk(
                content=current_chunk,
                source_file=source_file,
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=start_char + len(current_chunk)
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_long_section(
        self, 
        section: str, 
        source_file: str, 
        start_index: int,
        start_char: int
    ) -> List[DocumentChunk]:
        """Dividir sección muy larga en chunks manejables"""
        
        chunks = []
        words = section.split()
        current_chunk_words = []
        chunk_index = start_index
        current_start = start_char
        
        for word in words:
            test_chunk = ' '.join(current_chunk_words + [word])
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk_words.append(word)
            else:
                # Guardar chunk actual
                if current_chunk_words:
                    chunk_text = ' '.join(current_chunk_words)
                    chunk = self._create_chunk(
                        content=chunk_text,
                        source_file=source_file,
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(chunk_text)
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                    
                    # Aplicar overlap
                    if self.chunk_overlap > 0:
                        overlap_words = current_chunk_words[-self.chunk_overlap//10:]  # Aprox overlap en palabras
                        current_chunk_words = overlap_words + [word]
                    else:
                        current_chunk_words = [word]
                    
                    current_start += len(chunk_text) - (len(' '.join(overlap_words)) if self.chunk_overlap > 0 else 0)
                else:
                    current_chunk_words = [word]
        
        # Último chunk
        if current_chunk_words:
            chunk_text = ' '.join(current_chunk_words)
            chunk = self._create_chunk(
                content=chunk_text,
                source_file=source_file,
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(chunk_text)
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(
        self,
        content: str,
        source_file: str,
        chunk_index: int,
        start_char: int,
        end_char: int
    ) -> DocumentChunk:
        """Crear objeto DocumentChunk con metadatos"""
        
        file_path = Path(source_file)
        
        metadata = {
            'source_file': str(file_path.absolute()),
            'file_name': file_path.name,
            'file_extension': file_path.suffix.lower(),
            'file_size': file_path.stat().st_size if file_path.exists() else 0,
            'chunk_index': chunk_index,
            'chunk_size': len(content),
            'start_char': start_char,
            'end_char': end_char,
            'created_at': datetime.now().isoformat(),
            'processor_version': '1.0.0'
        }
        
        return DocumentChunk(
            id=str(uuid.uuid4()),
            content=content.strip(),
            metadata=metadata,
            source_file=source_file,
            chunk_index=chunk_index,
            start_char=start_char,
            end_char=end_char
        )
    
    def get_supported_extensions(self) -> List[str]:
        """Obtener lista de extensiones soportadas"""
        return list(self.processors.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """Verificar si un archivo es soportado"""
        extension = Path(file_path).suffix.lower()
        return extension in self.processors
    
    def get_processor_info(self) -> Dict[str, Any]:
        """Obtener información del procesador"""
        return {
            'supported_extensions': self.get_supported_extensions(),
            'chunk_size': self.chunk_size,
            'chunk_overlap': self.chunk_overlap,
            'separator': self.separator,
            'processors_available': {
                'pdf': PDF_AVAILABLE,
                'docx': DOCX_AVAILABLE,
                'pandas': PANDAS_AVAILABLE
            }
        }