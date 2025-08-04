"""
Servicio de ingesta multimodal para Prototipo_chatbot
Maneja documentos, APIs, web scraping y bases de datos
TFM Vicente Caruncho - Sistemas Inteligentes
"""

import os
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Generator, Tuple
from datetime import datetime
import concurrent.futures
from abc import ABC, abstractmethod

# Imports para procesamiento de documentos
import requests
from urllib.parse import urljoin, urlparse
import sqlite3
import json

# Imports locales
from app.core.config import get_ingestion_config
from app.core.logger import get_logger
from app.models import DocumentChunk, DocumentMetadata, IngestionJob, create_document_chunk

class DocumentProcessor(ABC):
    """Interfaz abstracta para procesadores de documentos"""
    
    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Verificar si puede procesar el archivo"""
        pass
    
    @abstractmethod
    def process(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Procesar archivo y retornar contenido y metadatos"""
        pass

class PDFProcessor(DocumentProcessor):
    """Procesador para archivos PDF"""
    
    def __init__(self):
        self.logger = get_logger("pdf_processor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    def process(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Procesar archivo PDF"""
        try:
            import PyPDF2
            
            content_parts = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content_parts.append(f"[Página {page_num + 1}]\n{text}")
                    except Exception as e:
                        self.logger.warning("Error extrayendo página PDF",
                                          file=str(file_path),
                                          page=page_num + 1,
                                          error=str(e))
            
            content = "\n\n".join(content_parts)
            
            # Crear metadatos
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                source_path=str(file_path),
                source_type='document',
                file_type='.pdf',
                size_bytes=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                processed_at=datetime.now(),
                checksum=self._calculate_checksum(file_path),
                page_count=page_count,
                word_count=len(content.split()) if content else 0
            )
            
            self.logger.info("PDF procesado exitosamente",
                           file=str(file_path),
                           pages=page_count,
                           content_length=len(content))
            
            return content, metadata
            
        except ImportError:
            raise ImportError("PyPDF2 no está instalado. Instala con: pip install PyPDF2")
        except Exception as e:
            self.logger.error("Error procesando PDF", file=str(file_path), error=str(e))
            raise
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calcular checksum MD5 del archivo"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

class DOCXProcessor(DocumentProcessor):
    """Procesador para archivos DOCX"""
    
    def __init__(self):
        self.logger = get_logger("docx_processor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.docx'
    
    def process(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Procesar archivo DOCX"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extraer texto de párrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            content = "\n\n".join(paragraphs)
            
            # Crear metadatos
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                source_path=str(file_path),
                source_type='document',
                file_type='.docx',
                size_bytes=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                processed_at=datetime.now(),
                checksum=self._calculate_checksum(file_path),
                word_count=len(content.split()) if content else 0
            )
            
            self.logger.info("DOCX procesado exitosamente",
                           file=str(file_path),
                           paragraphs=len(paragraphs),
                           content_length=len(content))
            
            return content, metadata
            
        except ImportError:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
        except Exception as e:
            self.logger.error("Error procesando DOCX", file=str(file_path), error=str(e))
            raise
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calcular checksum MD5 del archivo"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

class TXTProcessor(DocumentProcessor):
    """Procesador para archivos de texto plano"""
    
    def __init__(self):
        self.logger = get_logger("txt_processor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.md', '.rtf']
    
    def process(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Procesar archivo de texto"""
        try:
            # Detectar encoding
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            # Crear metadatos
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                source_path=str(file_path),
                source_type='document',
                file_type=file_path.suffix.lower(),
                size_bytes=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                processed_at=datetime.now(),
                checksum=self._calculate_checksum(file_path),
                encoding=encoding,
                word_count=len(content.split()) if content else 0
            )
            
            self.logger.info("Archivo de texto procesado exitosamente",
                           file=str(file_path),
                           encoding=encoding,
                           content_length=len(content))
            
            return content, metadata
            
        except Exception as e:
            self.logger.error("Error procesando archivo de texto",
                            file=str(file_path), error=str(e))
            raise
    
    def _detect_encoding(self, file_path: Path) -> str:
        """Detectar encoding del archivo"""
        try:
            # Intentar UTF-8 primero
            with open(file_path, 'r', encoding='utf-8') as f:
                f.read()
            return 'utf-8'
        except UnicodeDecodeError:
            try:
                # Intentar latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    f.read()
                return 'latin-1'
            except:
                # Fallback a cp1252 (Windows)
                return 'cp1252'
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calcular checksum MD5 del archivo"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

class ExcelProcessor(DocumentProcessor):
    """Procesador para archivos Excel"""
    
    def __init__(self):
        self.logger = get_logger("excel_processor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.xlsx', '.xls', '.csv']
    
    def process(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Procesar archivo Excel/CSV"""
        try:
            import pandas as pd
            
            # Leer archivo según su tipo
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path, encoding='utf-8')
            else:
                df = pd.read_excel(file_path)
            
            # Convertir a texto estructurado
            content_parts = []
            
            # Añadir información de columnas
            content_parts.append(f"Columnas: {', '.join(df.columns.tolist())}")
            content_parts.append(f"Número de filas: {len(df)}")
            content_parts.append("")
            
            # Convertir primeras 100 filas a texto
            for idx, row in df.head(100).iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                content_parts.append(f"Fila {idx + 1}: {row_text}")
            
            if len(df) > 100:
                content_parts.append(f"\n... ({len(df) - 100} filas adicionales no mostradas)")
            
            content = "\n".join(content_parts)
            
            # Crear metadatos
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                source_path=str(file_path),
                source_type='spreadsheet',
                file_type=file_path.suffix.lower(),
                size_bytes=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                processed_at=datetime.now(),
                checksum=self._calculate_checksum(file_path),
                word_count=len(content.split()) if content else 0
            )
            
            self.logger.info("Archivo Excel/CSV procesado exitosamente",
                           file=str(file_path),
                           rows=len(df),
                           columns=len(df.columns),
                           content_length=len(content))
            
            return content, metadata
            
        except ImportError:
            raise ImportError("pandas y openpyxl no están instalados. Instala con: pip install pandas openpyxl")
        except Exception as e:
            self.logger.error("Error procesando archivo Excel/CSV",
                            file=str(file_path), error=str(e))
            raise
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calcular checksum MD5 del archivo"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

class WebScraper:
    """Scraper para contenido web"""
    
    def __init__(self):
        self.logger = get_logger("web_scraper")
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Prototipo_chatbot/1.0 (TFM; Educational Use)'
        })
    
    def scrape_url(self, url: str) -> Tuple[str, DocumentMetadata]:
        """Scraper contenido de una URL"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Extraer contenido con BeautifulSoup
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remover scripts, estilos y otros elementos no deseados
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                tag.decompose()
            
            # Extraer título
            title = ""
            if soup.title:
                title = soup.title.get_text().strip()
            
            # Extraer contenido principal
            content_parts = []
            
            # Buscar contenido en elementos principales
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            if main_content:
                content_parts.append(main_content.get_text())
            else:
                # Fallback: extraer todo el texto del body
                if soup.body:
                    content_parts.append(soup.body.get_text())
            
            content = "\n".join(content_parts)
            
            # Limpiar texto
            content = self._clean_text(content)
            
            # Crear metadatos
            metadata = DocumentMetadata(
                source_path=url,
                source_type='web',
                file_type='.html',
                size_bytes=len(response.content),
                created_at=datetime.now(),
                processed_at=datetime.now(),
                checksum=hashlib.md5(response.content).hexdigest(),
                url=url,
                title=title,
                word_count=len(content.split()) if content else 0
            )
            
            self.logger.info("URL scrapeada exitosamente",
                           url=url,
                           title=title,
                           content_length=len(content))
            
            return content, metadata
            
        except ImportError:
            raise ImportError("beautifulsoup4 no está instalado. Instala con: pip install beautifulsoup4")
        except requests.exceptions.RequestException as e:
            self.logger.error("Error scrapeando URL", url=url, error=str(e))
            raise
        except Exception as e:
            self.logger.error("Error procesando contenido web", url=url, error=str(e))
            raise
    
    def _clean_text(self, text: str) -> str:
        """Limpiar texto extraído"""
        import re
        
        # Remover espacios en blanco excesivos
        text = re.sub(r'\s+', ' ', text)
        
        # Remover líneas vacías múltiples
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()

class TextChunker:
    """Chunker para dividir texto en fragmentos"""
    
    def __init__(self):
        self.config = get_ingestion_config()
        self.logger = get_logger("text_chunker")
    
    def chunk_text(self, content: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Dividir texto en chunks"""
        if not content.strip():
            return []
        
        chunks = []
        chunk_size = self.config.chunk_size
        chunk_overlap = self.config.chunk_overlap
        
        # Dividir en chunks con solapamiento
        start = 0
        chunk_index = 0
        
        while start < len(content):
            end = start + chunk_size
            
            # Ajustar final para no cortar palabras
            if end < len(content):
                # Buscar el final de la palabra más cercano
                while end > start and content[end] not in [' ', '\n', '\t', '.', '!', '?']:
                    end -= 1
                
                # Si no encontramos un buen punto de corte, usar el tamaño original
                if end == start:
                    end = start + chunk_size
            
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                try:
                    chunk = create_document_chunk(
                        content=chunk_content,
                        metadata=metadata,
                        chunk_index=chunk_index,
                        start_char=start
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                except Exception as e:
                    self.logger.warning("Error creando chunk",
                                      chunk_index=chunk_index,
                                      error=str(e))
            
            # Calcular siguiente posición con solapamiento
            start = end - chunk_overlap
            if start < 0:
                start = end
        
        self.logger.debug("Texto dividido en chunks",
                         total_chunks=len(chunks),
                         content_length=len(content),
                         chunk_size=chunk_size,
                         chunk_overlap=chunk_overlap)
        
        return chunks

class IngestionService:
    """Servicio principal de ingesta"""
    
    def __init__(self):
        self.logger = get_logger("ingestion_service")
        self.config = get_ingestion_config()
        
        # Inicializar procesadores
        self.processors = [
            PDFProcessor(),
            DOCXProcessor(),
            TXTProcessor(),
            ExcelProcessor()
        ]
        
        self.web_scraper = WebScraper()
        self.text_chunker = TextChunker()
        
        # Jobs activos
        self.active_jobs = {}
        
        self.logger.info("Servicio de ingesta inicializado",
                        processors=len(self.processors))
    
    def create_ingestion_job(self, job_type: str, sources: List[str],
                           config: Dict[str, Any] = None) -> IngestionJob:
        """Crear un trabajo de ingesta"""
        from app.models import create_ingestion_job
        
        job = create_ingestion_job(job_type, sources, config or {})
        self.active_jobs[job.id] = job
        
        self.logger.info("Trabajo de ingesta creado",
                        job_id=job.id,
                        job_type=job_type,
                        sources_count=len(sources))
        
        return job
    
    def process_documents(self, file_paths: List[str], 
                         job_id: Optional[str] = None) -> List[DocumentChunk]:
        """Procesar múltiples documentos"""
        all_chunks = []
        
        # Obtener job si se proporciona
        job = self.active_jobs.get(job_id) if job_id else None
        
        if job:
            job.start()
            job.total_items = len(file_paths)
        
        # Procesar archivos en paralelo si está configurado
        if self.config.parallel_processing and len(file_paths) > 1:
            with concurrent.futures.ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
                future_to_path = {
                    executor.submit(self._process_single_document, path): path 
                    for path in file_paths
                }
                
                for i, future in enumerate(concurrent.futures.as_completed(future_to_path)):
                    file_path = future_to_path[future]
                    
                    try:
                        chunks = future.result()
                        all_chunks.extend(chunks)
                        
                        if job:
                            job.update_progress(file_path, i + 1)
                            job.created_chunks += len(chunks)
                            
                    except Exception as e:
                        self.logger.error("Error procesando documento",
                                        file=file_path, error=str(e))
                        if job:
                            job.failed_items += 1
                            job.error_messages.append(f"{file_path}: {str(e)}")
        else:
            # Procesamiento secuencial
            for i, file_path in enumerate(file_paths):
                try:
                    chunks = self._process_single_document(file_path)
                    all_chunks.extend(chunks)
                    
                    if job:
                        job.update_progress(file_path, i + 1)
                        job.created_chunks += len(chunks)
                        
                except Exception as e:
                    self.logger.error("Error procesando documento",
                                    file=file_path, error=str(e))
                    if job:
                        job.failed_items += 1
                        job.error_messages.append(f"{file_path}: {str(e)}")
        
        if job:
            job.complete()
            
        self.logger.info("Procesamiento de documentos completado",
                        files_processed=len(file_paths),
                        chunks_created=len(all_chunks),
                        job_id=job_id)
        
        return all_chunks
    
    def _process_single_document(self, file_path: str) -> List[DocumentChunk]:
        """Procesar un solo documento"""
        path = Path(file_path)
        
        # Verificar que el archivo existe
        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        # Verificar tamaño del archivo
        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.max_file_size_mb:
            raise ValueError(f"Archivo demasiado grande: {file_size_mb:.2f}MB > {self.config.max_file_size_mb}MB")
        
        # Encontrar procesador apropiado
        processor = None
        for proc in self.processors:
            if proc.can_process(path):
                processor = proc
                break
        
        if not processor:
            raise ValueError(f"No hay procesador disponible para: {path.suffix}")
        
        # Procesar documento
        content, metadata = processor.process(path)
        
        # Dividir en chunks
        chunks = self.text_chunker.chunk_text(content, metadata)
        
        return chunks
    
    def process_urls(self, urls: List[str], job_id: Optional[str] = None) -> List[DocumentChunk]:
        """Procesar múltiples URLs"""
        all_chunks = []
        
        job = self.active_jobs.get(job_id) if job_id else None
        
        if job:
            job.start()
            job.total_items = len(urls)
        
        for i, url in enumerate(urls):
            try:
                content, metadata = self.web_scraper.scrape_url(url)
                chunks = self.text_chunker.chunk_text(content, metadata)
                all_chunks.extend(chunks)
                
                if job:
                    job.update_progress(url, i + 1)
                    job.created_chunks += len(chunks)
                    
            except Exception as e:
                self.logger.error("Error procesando URL", url=url, error=str(e))
                if job:
                    job.failed_items += 1
                    job.error_messages.append(f"{url}: {str(e)}")
        
        if job:
            job.complete()
        
        self.logger.info("Procesamiento de URLs completado",
                        urls_processed=len(urls),
                        chunks_created=len(all_chunks),
                        job_id=job_id)
        
        return all_chunks
    
    def get_job_status(self, job_id: str) -> Optional[IngestionJob]:
        """Obtener estado de un trabajo"""
        return self.active_jobs.get(job_id)
    
    def get_active_jobs(self) -> List[IngestionJob]:
        """Obtener todos los trabajos activos"""
        return list(self.active_jobs.values())
    
    def cancel_job(self, job_id: str) -> bool:
        """Cancelar un trabajo"""
        job = self.active_jobs.get(job_id)
        if job and job.status in ['pending', 'running']:
            job.status = 'cancelled'
            job.completed_at = datetime.now()
            self.logger.info("Trabajo cancelado", job_id=job_id)
            return True
        return False
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Obtener formatos soportados"""
        return {
            'documents': self.config.document_formats,
            'spreadsheets': self.config.spreadsheet_formats,
            'web': self.config.web_formats
        }
    
    def get_service_stats(self) -> Dict[str, Any]:
        """Obtener estadísticas del servicio"""
        active_jobs = len([j for j in self.active_jobs.values() if j.status == 'running'])
        completed_jobs = len([j for j in self.active_jobs.values() if j.status == 'completed'])
        failed_jobs = len([j for j in self.active_jobs.values() if j.status == 'failed'])
        
        return {
            'processors_available': len(self.processors),
            'supported_formats': self.get_supported_formats(),
            'max_file_size_mb': self.config.max_file_size_mb,
            'parallel_processing': self.config.parallel_processing,
            'max_workers': self.config.max_workers,
            'chunk_size': self.config.chunk_size,
            'chunk_overlap': self.config.chunk_overlap,
            'jobs_active': active_jobs,
            'jobs_completed': completed_jobs,
            'jobs_failed': failed_jobs,
            'total_jobs': len(self.active_jobs)
        }

# Instancia global del servicio de ingesta
ingestion_service = IngestionService()

# Funciones de conveniencia
def process_documents(file_paths: List[str]) -> List[DocumentChunk]:
    """Función de conveniencia para procesar documentos"""
    return ingestion_service.process_documents(file_paths)

def process_urls(urls: List[str]) -> List[DocumentChunk]:
    """Función de conveniencia para procesar URLs"""
    return ingestion_service.process_urls(urls)

def get_ingestion_stats() -> Dict[str, Any]:
    """Función de conveniencia para obtener estadísticas"""
    return ingestion_service.get_service_stats()